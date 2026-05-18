"""Generate the sysadmin questionnaire as a .docx file.

Usage:
    py -3.10 scripts/build_sysadmin_doc.py

Output: docs/sysadmin-questions.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


def add_heading(doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if level == 0:
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def add_bullet(doc, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_question(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    # Empty line for the answer
    ans = doc.add_paragraph()
    ans.paragraph_format.left_indent = Cm(1.0)
    r = ans.add_run("Ответ: ____________________________________________________________")
    r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = val


def build() -> Path:
    doc = Document()

    # --- Title ---
    add_heading(doc, "JiraAnalysis — вопросы системному администратору", level=0)
    add_para(
        doc,
        "Развёртывание корпоративного сервиса через GitLab CI с автоматической сборкой образа из исходников. "
        "Документ заполняется один раз; после ответов разработчик финализирует конфиги и передаёт пакет к развёртыванию.",
        italic=True,
    )
    doc.add_paragraph()

    # --- Что уже готово ---
    add_heading(doc, "Что готово со стороны разработчика", level=1)
    add_bullet(doc, "Dockerfile (multi-stage: node-frontend → python-deps + e5-модель → slim runtime)")
    add_bullet(doc, "docker-compose.yml для prod и staging (postgres + backend + nginx)")
    add_bullet(doc, "nginx.conf с поддержкой SSE-эндпоинтов")
    add_bullet(doc, "Шаблон .env для всех секретов")
    add_bullet(doc, "Скрипт миграции SQLite → PostgreSQL")
    add_bullet(doc, "Healthcheck-эндпоинты /health (liveness) и /health/ready (DB readiness)")
    add_bullet(doc, "Alembic-миграции для PostgreSQL")
    doc.add_paragraph()

    # ============== Section 1: GitLab CI ==============
    add_heading(doc, "1. GitLab CI и сборка образа", level=1)
    add_para(doc, "Подтверждённая схема: GitLab pull-ит исходники из GitHub, валидирует, собирает образ, разворачивает на ВМ.")
    doc.add_paragraph()
    add_question(doc, "URL вашего GitLab и path репозитория?")
    add_question(doc, "Что триггерит сборку — push в main / тег vX.Y.Z / merge MR / ручной запуск?")
    add_question(doc, "Какая ветка/тег идёт на prod, какая на staging?")
    add_question(doc, "Зеркало GitHub→GitLab двустороннее или только pull? "
                      "Достаточно ли push в GitHub чтобы изменения появились в GitLab автоматически?")
    add_question(doc, "Кто отвечает за добавление новых веток/тегов в синхронизацию (если нужно)?")
    doc.add_paragraph()

    # ============== Section 2: Container registry ==============
    add_heading(doc, "2. Container Registry", level=1)
    add_question(doc, "URL внутреннего Docker registry (Harbor / Nexus / GitLab Container Registry)?")
    add_question(doc, "Под каким именем будет публиковаться образ? Например: registry.corp/jira-analytics/backend")
    add_question(doc, "Как ВМ авторизуется в registry — service-account, imagePullSecret, anonymous? Где хранятся креды?")
    add_question(doc, "Сколько последних тегов хранится в registry? Нужен ли retention policy для отката?")
    doc.add_paragraph()

    # ============== Section 3: Deploy mechanism ==============
    add_heading(doc, "3. Механизм развёртывания на ВМ", level=1)
    add_question(doc, "Кто запускает docker compose pull && docker compose up -d на ВМ — "
                      "GitLab runner через SSH/ansible, k8s-оператор, или ваша рука?")
    add_question(doc, "Где живёт docker-compose.yml на ВМ — в репо (синкается из git), или вы кладёте отдельно?")
    add_question(doc, "Где хранится .env — в GitLab CI variables (инжектится при деплое) или persistent-файл на ВМ?")
    add_question(doc, "Pipeline ждёт ответ от GET /health/ready (HTTP 200) перед тем как считать deploy успешным? "
                      "Если нет — кто валидирует работоспособность нового образа?")
    add_question(doc, "Какой ожидаемый downtime на один deploy? (recreate контейнера ~30-60 секунд)")
    doc.add_paragraph()

    # ============== Section 4: Build environment ==============
    add_heading(doc, "4. Окружение GitLab runner для сборки образа", level=1)
    add_para(doc, "Сборка требует исходящий доступ. Если что-то закрыто — нужны корпоративные mirror-ы или pre-built base image:")
    add_table(
        doc,
        headers=["Цель", "Хост", "Зачем"],
        rows=[
            ["npm", "registry.npmjs.org", "npm ci для фронта (~700MB зависимостей)"],
            ["pip", "pypi.org, files.pythonhosted.org", "pip install backend-зависимостей"],
            ["torch", "download.pytorch.org", "CPU-сборка torch ~200MB"],
            ["HuggingFace", "huggingface.co", "Скачивание модели multilingual-e5-base ~470MB"],
            ["apt", "deb.debian.org", "apt install build-essential, libpq-dev, curl, tini"],
        ],
    )
    doc.add_paragraph()
    add_question(doc, "Что из этого закрыто? Какие корп. mirror-ы есть на замену?")
    add_question(doc, "Если HuggingFace недоступен — можно ли разработчику собрать base image отдельно "
                      "(с torch + моделью внутри), загрузить во внутренний registry, и наш Dockerfile будет наследоваться от него?")
    add_question(doc, "Лимит RAM/CPU/disk для GitLab runner на одну сборку? (Build тяжёлый: ~5GB образ, ~10 минут.)")
    add_question(doc, "Версия Docker и наличие buildx на runner? (для multi-stage build с кэшем)")
    doc.add_paragraph()

    # ============== Section 5: Database migrations ==============
    add_heading(doc, "5. Применение миграций БД", level=1)
    add_para(doc, "alembic upgrade head применяет схему Postgres. На каждом релизе могут быть новые миграции.")
    doc.add_paragraph()
    add_question(doc, "Кто запускает alembic upgrade head — отдельный GitLab job до docker compose up, или backend на старте? "
                      "Рекомендуется отдельный job — упавшая миграция остановит pipeline до того как сломанная схема пойдёт в prod.")
    add_question(doc, "Если миграция падает — pipeline откатывается полностью, или образ всё равно деплоится с битой схемой?")
    add_question(doc, "Есть ли механизм блокировки одновременных миграций (только один pod применяет схему)?")
    doc.add_paragraph()

    # ============== Section 6: Rollback ==============
    add_heading(doc, "6. Откат при битом деплое", level=1)
    add_question(doc, "GitLab держит предыдущий образ под отдельным тегом для быстрого rollback?")
    add_question(doc, "Команда rollback — re-run прошлого pipeline / кнопка в UI / ручной change tag?")
    add_question(doc, "Сколько времени занимает rollback от момента обнаружения проблемы до восстановления?")
    add_question(doc, "Если миграция уже применилась к Postgres и пошёл rollback образа — старый код работает с новой схемой? "
                      "(Мы стараемся писать forward-compatible миграции, но нужен план Б на ЧП.)")
    doc.add_paragraph()

    # ============== Section 7: Secrets ==============
    add_heading(doc, "7. Хранение секретов", level=1)
    add_para(doc, "Сервис требует:")
    add_table(
        doc,
        headers=["Переменная", "Тип", "Lifecycle"],
        rows=[
            ["DB_PASSWORD", "Пароль Postgres", "Генерируется один раз, постоянный"],
            ["JWT_SECRET_KEY", "32-байтный ключ", "Один раз. Ротация = разлогин всех пользователей"],
            ["GEMINI_API_KEY", "API-ключ Google", "Ротация по запросу разработчика"],
            ["OPENROUTER_API_KEY", "API-ключ OpenRouter", "Аналогично"],
            ["ADMIN_PASSWORD", "Временный", "Удаляется после первого входа"],
        ],
    )
    doc.add_paragraph()
    add_question(doc, "Где живут эти секреты — GitLab CI variables / HashiCorp Vault / persistent .env на ВМ?")
    add_question(doc, "Кто генерирует JWT_SECRET_KEY и DB_PASSWORD при первой установке — вы или разработчик передаёт?")
    add_question(doc, "Куда сохраняем backup (если файл .env потеряется, JWT-ключ восстановить нельзя)?")
    add_question(doc, "LLM-ключи (Gemini / OpenRouter) — кто их выдаёт и куда вписывать?")
    doc.add_paragraph()

    # ============== Section 8: Server / network ==============
    add_heading(doc, "8. Сервер и сеть", level=1)
    add_question(doc, "Характеристики ВМ (vCPU / RAM / диск). Рекомендованный минимум: 4 vCPU / 16 GB RAM / 100 GB SSD.")
    add_question(doc, "ОС и её версия?")
    add_question(doc, "Внутренний домен сервиса (например jira-analytics.itgri.local). "
                      "Это настоящее DNS-имя или нужно добавлять?")
    add_question(doc, "Будет ли staging доступен по отдельному hostname или только по IP+порт?")
    add_question(doc, "Какие порты сервера открыты клиентам в корп. сети? Ожидаем 443 (HTTPS) обязательно, 80 (HTTP→HTTPS редирект) желательно.")
    add_question(doc, "Доступ только из офисной подсети или через VPN тоже?")
    doc.add_paragraph()

    # ============== Section 9: Outbound traffic from server ==============
    add_heading(doc, "9. Исходящий трафик с ВМ в runtime", level=1)
    add_para(doc, "В отличие от build-окружения (Section 4), production VM требует доступ только к этим адресам:")
    add_table(
        doc,
        headers=["Цель", "Хост", "Зачем"],
        rows=[
            ["Jira Cloud", "itgri.atlassian.net (HTTPS:443)", "Синхронизация задач/ворклогов (постоянно)"],
            ["Gemini API", "generativelanguage.googleapis.com (HTTPS:443)", "AI-саммари проектов (опционально)"],
            ["OpenRouter", "openrouter.ai (HTTPS:443)", "Резервный LLM-провайдер (опционально)"],
            ["RU календарь", "data.gov.ru или мирроры (HTTPS:443)", "Синхронизация производственного календаря"],
        ],
    )
    doc.add_paragraph()
    add_question(doc, "Что из этого открыто? Если корп. прокси — какой URL и порт?")
    add_question(doc, "Если что-то полностью закрыто — какие альтернативы (mirror, отказ от фичи)?")
    doc.add_paragraph()

    # ============== Section 10: TLS ==============
    add_heading(doc, "10. TLS и обратный прокси", level=1)
    add_para(doc, "Варианты: (а) корп. CA выдаёт сертификат, мы кладём в nginx; (б) внешний balancer терминирует TLS, backend по HTTP.")
    doc.add_paragraph()
    add_question(doc, "Какой вариант — собственный nginx с сертификатом от корп. CA, или внешний balancer?")
    add_question(doc, "Если корп. CA: какой срок действия и кто отвечает за продление сертификата?")
    add_question(doc, "Если внешний balancer: поддерживаются ли SSE через ваш балансер? "
                      "(длинноживущие HTTP-соединения, timeout ≥ 1 час, no buffering)")
    add_question(doc, "Если на сервере уже стоит общий nginx/traefik — можем выкинуть наш nginx-контейнер и сидеть за вашим прокси?")
    doc.add_paragraph()

    # ============== Section 11: Backups ==============
    add_heading(doc, "11. Резервные копии", level=1)
    add_question(doc, "Частота VM-снапшотов и срок хранения?")
    add_question(doc, "Тестировался ли restore? Сколько времени занимает откат?")
    add_question(doc, "Возможен ли дополнительный ручной снапшот перед каждым релизом (как safety-net)?")
    add_question(doc, "Нужен ли отдельный pg_dump-cron в дополнение к VM-снапшотам? Куда складывать дампы?")
    doc.add_paragraph()

    # ============== Section 12: Initial cutover ==============
    add_heading(doc, "12. Перенос существующей БД (однократно)", level=1)
    add_para(doc, "У разработчика локально SQLite ~457MB с данными. После запуска инфраструктуры это нужно перенести в production Postgres.")
    doc.add_paragraph()
    add_question(doc, "Как передать SQLite-файл с моего ПК на ВМ? Варианты: SFTP, внутренний файлообменник, временный SSH, корп. S3.")
    add_question(doc, "Кто запустит скрипт миграции scripts/migrate_to_postgres.py на ВМ — вы по runbook-у?")
    add_question(doc, "После миграции файл-снимок нужно удалить — кто и когда?")
    doc.add_paragraph()

    # ============== Section 13: Monitoring ==============
    add_heading(doc, "13. Мониторинг", level=1)
    add_question(doc, "Есть ли корп. uptime-инструмент (Zabbix, Prometheus, Grafana)? "
                      "Endpoint для проверки: GET https://<домен>/health → 200")
    add_question(doc, "Если есть централизованная система логов (ELK, Loki, Splunk) — нужно ли настроить Docker logging driver?")
    add_question(doc, "Куда слать алерты — email / Telegram / корп. чат?")
    doc.add_paragraph()

    # ============== Section 14: Roles / access ==============
    add_heading(doc, "14. Доступы и регламент", level=1)
    add_question(doc, "Кто второй человек на случай вашего отсутствия (отпуск, болезнь)?")
    add_question(doc, "Часовой пояс — подтверждаете Europe/Moscow? (производственный календарь считает по Москве)")
    add_question(doc, "Окно деплоя — конкретное время или по запросу?")
    add_question(doc, "Куда писать при ЧП (502 / падение контейнера) в нерабочее время?")
    doc.add_paragraph()

    # ============== Final checklist ==============
    add_heading(doc, "Чек-лист готовности к развёртыванию", level=1)
    add_para(doc, "Минимум для запуска инфраструктуры (порядок не критичен):")
    add_bullet(doc, "Раздел 1.1 — URL GitLab и project path")
    add_bullet(doc, "Раздел 1.2 — Триггер сборки определён")
    add_bullet(doc, "Раздел 2.1-2.2 — Registry URL и образ-path")
    add_bullet(doc, "Раздел 3.1 — Механизм deploy на ВМ определён")
    add_bullet(doc, "Раздел 4.1 — Build-окружение имеет нужный outbound (или есть mirror/base image)")
    add_bullet(doc, "Раздел 7.1 — Где живут секреты")
    add_bullet(doc, "Раздел 8.1-8.3 — Сервер выделен, домен прописан")
    add_bullet(doc, "Раздел 10.1 — Источник TLS-сертификата")
    add_bullet(doc, "Раздел 12.1 — Канал передачи SQLite-снимка")
    doc.add_paragraph()
    add_para(
        doc,
        "Когда минимум закрыт — разработчик финализирует docker-compose.yml (registry-path), "
        ".env-шаблон, удалит obsolete GitHub Actions release pipeline и передаст пакет к развёртыванию.",
        italic=True,
    )

    out = Path("docs/sysadmin-questions.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print(f"Wrote: {path.resolve()}")
